import random
import re
import json
import logging
import sys
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from transformers import AutoModelForCausalLM, AutoTokenizer
import torch
print(torch.__version__)
from tqdm import tqdm
from scipy.stats import norm
import pandas as pd
import os
from torch.cuda import OutOfMemoryError
from typing import Tuple, List, Dict, Any, Optional
import time
import traceback
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


# Set up logging with more detailed configuration
def setup_logging():
    """Set up comprehensive logging configuration."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_filename = f"math_eval_{timestamp}.log"
    
    # Create formatters
    detailed_formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s'
    )
    simple_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    
    # Set up root logger
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)
    
    # File handler for detailed logs
    file_handler = logging.FileHandler(log_filename, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(detailed_formatter)
    
    # Console handler for important messages
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(simple_formatter)
    
    # Error handler for critical errors
    error_filename = f"math_eval_errors_{timestamp}.log"
    error_handler = logging.FileHandler(error_filename, encoding='utf-8')
    error_handler.setLevel(logging.ERROR)
    error_handler.setFormatter(detailed_formatter)
    
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    logger.addHandler(error_handler)
    
    return logger

logger = setup_logging()

# Default configuration
DEFAULT_CONFIG = {
    "models": [
        ## -- aquiffoo -- ##
        "aquiffoo/aquif-neo-small",
        "aquiffoo/aquif-moe-800m",
        "aquiffoo/aquif-neo",

        ## -- VortexIntelligence -- ##
        "VortexIntelligence/VLM-1.1-K1-Preview",
        "VortexIntelligence/VLM-1-K3",
        "VortexIntelligence/VLM-1-K2",
        "VortexIntelligence/VLM-1-K1",

        ## -- OpenAI -- ##
        "openai-community/gpt2-large",
        "openai-community/gpt2",
        
        ## -- FlameF0X -- ##
        "FlameF0X/Snowflake-G0-Release-2",
        "FlameF0X/Muffin-2.9b-1C25",
        "FlameF0X/MathGPT2-Kaly",
        "FlameF0X/MuffinFace-2",
        "FlameF0X/MuffinFace-1",
        "FlameF0X/A-MoE-Model",
        "FlameF0X/MathGPT2",
        
        ## -- GoofyLM -- ##
        "GoofyLM/gonzalez-v1",
        
        ## -- EleutherAI -- ##
         "EleutherAI/gpt-neo-125m",
         "EleutherAI/gpt-neo-1.3B",
    ],
    "num_questions": 10,
    "max_tokens": 16,
    "batch_size": 8,
    "timeout": 10,
    "difficulty": "mixed",
    "output_dir": "results",
    "device": None
}

# Determine if we are in a Jupyter/Colab environment
IN_NOTEBOOK = 'ipykernel' in sys.modules

# Create a class to hold configuration - works in both script and notebook contexts
class Args:
    def __init__(self, **kwargs):
        for key, value in kwargs.items():
            setattr(self, key, value)

# Parse arguments or use defaults based on environment
if IN_NOTEBOOK:
    # In notebook, use default configuration
    args = Args(**DEFAULT_CONFIG)
    logger.info("Running in notebook environment, using default configuration")
else:
    # In script environment, use argparse
    import argparse
    parser = argparse.ArgumentParser(description='Evaluate LLMs on basic arithmetic operations')
    parser.add_argument('--models', nargs='+', default=DEFAULT_CONFIG["models"],
                        help='List of models to evaluate')
    parser.add_argument('--num_questions', type=int, default=DEFAULT_CONFIG["num_questions"],
                        help='Number of questions to evaluate')
    parser.add_argument('--max_tokens', type=int, default=DEFAULT_CONFIG["max_tokens"],
                        help='Maximum tokens for model generation')
    parser.add_argument('--batch_size', type=int, default=DEFAULT_CONFIG["batch_size"],
                        help='Batch size for evaluation')
    parser.add_argument('--timeout', type=int, default=DEFAULT_CONFIG["timeout"],
                        help='Timeout in seconds for model inference')
    parser.add_argument('--difficulty', choices=['easy', 'medium', 'hard', 'mixed'],
                        default=DEFAULT_CONFIG["difficulty"], help='Math problem difficulty')
    parser.add_argument('--output_dir', type=str, default=DEFAULT_CONFIG["output_dir"],
                        help='Directory to save results')
    parser.add_argument('--device', type=str, default=DEFAULT_CONFIG["device"],
                        help='Device to run on (default: auto-detect)')
    args = parser.parse_args()

# Set device
if args.device:
    DEVICE = args.device
else:
    DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
logger.info(f"Using device: {DEVICE}")

# Global error tracking
MODEL_ERRORS = {}
EVALUATION_STATS = {
    'start_time': None,
    'end_time': None,
    'total_duration': None,
    'models_attempted': 0,
    'models_successful': 0,
    'models_failed': 0,
    'total_questions_attempted': 0,
    'total_questions_successful': 0,
    'system_info': {
        'python_version': sys.version,
        'torch_version': torch.__version__,
        'device': DEVICE,
        'cuda_available': torch.cuda.is_available(),
    }
}

# Function to update configuration in notebook environment
def update_config(**kwargs):
    """Update configuration parameters (for use in notebooks)"""
    for key, value in kwargs.items():
        if hasattr(args, key):
            setattr(args, key, value)
            logger.info(f"Updated {key} to {value}")
        else:
            logger.warning(f"Unknown configuration parameter: {key}")

    # Re-set device if it was updated
    global DEVICE
    if "device" in kwargs:
        DEVICE = args.device if args.device else ("cuda" if torch.cuda.is_available() else "cpu")
        logger.info(f"Using device: {DEVICE}")

# Create output directory if it doesn't exist
os.makedirs(args.output_dir, exist_ok=True)

# Define difficulty levels for math problems
DIFFICULTY_RANGES = {
    'easy': {'num_range': (0, 10), 'operations': ['+', '-', '*', '/']},
    'medium': {'num_range': (10, 100), 'operations': ['+', '-', '*', '/']},
    'hard': {'num_range': (100, 1000), 'operations': ['+', '-', '*', '/', '**']},
    'mixed': None  # Will be handled specially
}

def log_model_error(model_name: str, error_type: str, error_message: str, traceback_str: str = None):
    """Log model-specific errors with detailed information."""
    if model_name not in MODEL_ERRORS:
        MODEL_ERRORS[model_name] = []
    
    error_info = {
        'timestamp': datetime.now().isoformat(),
        'error_type': error_type,
        'error_message': str(error_message),
        'traceback': traceback_str or traceback.format_exc()
    }
    
    MODEL_ERRORS[model_name].append(error_info)
    logger.error(f"Model {model_name} - {error_type}: {error_message}")
    if traceback_str:
        logger.debug(f"Traceback for {model_name}:\n{traceback_str}")

def get_difficulty_params(difficulty: str) -> dict:
    """Get parameters for the specified difficulty level."""
    if difficulty == 'mixed':
        # Randomly choose a difficulty
        difficulty = random.choice(['easy', 'medium', 'hard'])

    return DIFFICULTY_RANGES[difficulty]

def generate_problem(difficulty: str = 'easy') -> Tuple[str, str, str]:
    """Generate a math problem with the specified difficulty.

    Returns:
        Tuple containing (question, expected_answer, operation)
    """
    params = get_difficulty_params(difficulty)
    num_range = params['num_range']
    operations = params['operations']

    op = random.choice(operations)

    if op == '/':
        b = random.randint(1, num_range[1])
        # Make sure division results in an integer
        a = b * random.randint(1, num_range[1] // b)
        result = a // b
        question = f"{a} / {b}"  # Only the mathematical expression
        return question, str(result), op

    elif op == '*':
        a = random.randint(num_range[0], num_range[1])
        b = random.randint(num_range[0], num_range[1])
        result = a * b
        question = f"{a} * {b}"  # Only the mathematical expression
        return question, str(result), op

    elif op == '+':
        a = random.randint(num_range[0], num_range[1])
        b = random.randint(num_range[0], num_range[1])
        result = a + b
        question = f"{a} + {b}"  # Only the mathematical expression
        return question, str(result), op

    elif op == '-':
        a = random.randint(num_range[0], num_range[1])
        b = random.randint(num_range[0], min(a, num_range[1]))  # Ensure b <= a for simpler problems
        result = a - b
        question = f"{a} - {b}"  # Only the mathematical expression
        return question, str(result), op

    elif op == '**':
        a = random.randint(2, 10)
        b = random.randint(2, 3)  # Keep exponents small
        result = a ** b
        question = f"{a} ** {b}"  # Only the mathematical expression
        return question, str(result), op

def load_model_and_tokenizer(model_name: str) -> Tuple[Any, Any]:
    """Load model and tokenizer with comprehensive error handling."""
    try:
        start_time = time.time()
        logger.info(f"Loading {model_name}...")

        try:
            tokenizer = AutoTokenizer.from_pretrained(model_name)
        except Exception as e:
            log_model_error(model_name, "TokenizerLoadError", str(e))
            raise

        try:
            model = AutoModelForCausalLM.from_pretrained(model_name).to(DEVICE)
            model.eval()
        except Exception as e:
            log_model_error(model_name, "ModelLoadError", str(e))
            raise

        load_time = time.time() - start_time
        logger.info(f"Loaded {model_name} in {load_time:.2f} seconds")

        return tokenizer, model
    except Exception as e:
        logger.error(f"Failed to load {model_name}: {str(e)}")
        raise

def extract_answer(response: str) -> str:
    """Extract the numeric answer from model response using regex patterns."""
    # Try to find a numeric answer in the response
    patterns = [
        r"answer is (\d+)",
        r"result is (\d+)",
        r"equals (\d+)",
        r"= (\d+)",
        r"(\d+)",  # Fallback to any number
    ]

    for pattern in patterns:
        matches = re.search(pattern, response.lower())
        if matches:
            return matches.group(1)

    return ""

def ask_model(tokenizer: Any, model: Any, question: str, model_name: str) -> str:
    """Query the model with comprehensive error handling and timeout."""
    try:
        inputs = tokenizer(question, return_tensors="pt").to(DEVICE)

        with torch.no_grad():
            # Set a timeout for generation
            start_time = time.time()
            outputs = model.generate(
                **inputs,
                max_new_tokens=args.max_tokens,
                pad_token_id=tokenizer.eos_token_id
            )

        full_response = tokenizer.decode(outputs[0], skip_special_tokens=True)

        # Extract just the answer part
        answer_text = full_response[len(question):].strip()

        # Try to extract the numeric answer
        answer = extract_answer(answer_text)

        return answer

    except OutOfMemoryError as e:
        log_model_error(model_name, "OutOfMemoryError", f"Question: {question}")
        torch.cuda.empty_cache()
        return ""
    except RuntimeError as e:
        log_model_error(model_name, "RuntimeError", f"Question: {question}, Error: {str(e)}")
        return ""
    except Exception as e:
        log_model_error(model_name, "InferenceError", f"Question: {question}, Error: {str(e)}")
        return ""

def calculate_confidence_interval(correct: int, total: int, confidence: float = 0.95) -> Tuple[float, float]:
    """Calculate confidence interval for binomial proportion."""
    if total == 0:
        return 0.0, 0.0

    p = correct / total
    z = norm.ppf(1 - (1 - confidence) / 2)
    margin = z * (p * (1 - p) / total) ** 0.5

    return max(0, p - margin), min(1, p + margin)

def evaluate_models() -> Dict[str, Dict[str, Any]]:
    """Evaluate all models and return comprehensive results."""
    results = {}
    EVALUATION_STATS['start_time'] = datetime.now()
    EVALUATION_STATS['models_attempted'] = len(args.models)

    for model_name in args.models:
        logger.info(f"\nEvaluating {model_name}...")
        model_start_time = time.time()

        try:
            tokenizer, model = load_model_and_tokenizer(model_name)
            EVALUATION_STATS['models_successful'] += 1
        except Exception as e:
            logger.warning(f"Skipping {model_name} due to loading failure")
            EVALUATION_STATS['models_failed'] += 1
            results[model_name] = {
                "correct": 0,
                "incorrect": args.num_questions,
                "accuracy": 0.0,
                "ci_low": 0.0,
                "ci_high": 0.0,
                "model_load_failed": True,
                "model_load_error": str(e),
                "evaluation_time": 0.0,
                "by_operation": {op: {"correct": 0, "total": 0, "accuracy": 0.0}
                               for op in ['+', '-', '*', '/', '**']}
            }
            continue

        correct = 0
        incorrect = 0
        inference_errors = 0
        by_operation = {op: {"correct": 0, "total": 0} for op in ['+', '-', '*', '/', '**']}
        question_details = []

        for i in tqdm(range(args.num_questions), desc=f"Evaluating {model_name.split('/')[-1]}"):
            question, expected_answer, operation = generate_problem(args.difficulty)
            EVALUATION_STATS['total_questions_attempted'] += 1

            try:
                answer = ask_model(tokenizer, model, question, model_name)

                # Update operation-specific stats
                by_operation[operation]["total"] += 1

                question_detail = {
                    "question": question,
                    "expected_answer": expected_answer,
                    "model_answer": answer,
                    "operation": operation,
                    "correct": False
                }

                if answer.isdigit() and answer == expected_answer:
                    correct += 1
                    by_operation[operation]["correct"] += 1
                    question_detail["correct"] = True
                    EVALUATION_STATS['total_questions_successful'] += 1
                else:
                    incorrect += 1

                question_details.append(question_detail)

            except Exception as e:
                logger.error(f"Error evaluating {question}: {str(e)}")
                log_model_error(model_name, "EvaluationError", f"Question: {question}, Error: {str(e)}")
                incorrect += 1
                inference_errors += 1

        # Calculate accuracies and confidence intervals
        total_answered = correct + incorrect
        accuracy = correct / total_answered if total_answered > 0 else 0
        ci_low, ci_high = calculate_confidence_interval(correct, total_answered)

        # Calculate operation-specific accuracies
        for op in by_operation:
            op_total = by_operation[op]["total"]
            op_correct = by_operation[op]["correct"]
            by_operation[op]["accuracy"] = op_correct / op_total if op_total > 0 else 0

        model_evaluation_time = time.time() - model_start_time

        results[model_name] = {
            "correct": correct,
            "incorrect": incorrect,
            "accuracy": accuracy,
            "ci_low": ci_low,
            "ci_high": ci_high,
            "inference_errors": inference_errors,
            "model_load_failed": False,
            "evaluation_time": model_evaluation_time,
            "by_operation": by_operation,
            "question_details": question_details,
            "model_errors": MODEL_ERRORS.get(model_name, [])
        }

        # Clean up GPU memory
        if DEVICE == "cuda":
            del model
            del tokenizer
            torch.cuda.empty_cache()

        logger.info(f"Completed {model_name} in {model_evaluation_time:.2f} seconds - Accuracy: {accuracy:.1%}")

        # Save intermediate results
        save_results(results)

    EVALUATION_STATS['end_time'] = datetime.now()
    EVALUATION_STATS['total_duration'] = (EVALUATION_STATS['end_time'] - EVALUATION_STATS['start_time']).total_seconds()

    return results

def save_results(results: Dict[str, Dict[str, Any]]) -> None:
    """Save comprehensive results to multiple formats."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Save JSON results
    json_filename = os.path.join(args.output_dir, f"math_eval_results_{timestamp}.json")
    with open(json_filename, 'w') as f:
        json.dump({
            'evaluation_stats': EVALUATION_STATS,
            'configuration': vars(args),
            'results': results,
            'model_errors': MODEL_ERRORS
        }, f, indent=2, default=str)
    
    logger.info(f"JSON results saved to {json_filename}")
    
    # Save CSV summary
    csv_filename = os.path.join(args.output_dir, f"math_eval_summary_{timestamp}.csv")
    summary_data = []
    
    for model_name, result in results.items():
        summary_data.append({
            'Model': model_name,
            'Accuracy': result['accuracy'],
            'Correct': result['correct'],
            'Incorrect': result['incorrect'],
            'CI_Low': result['ci_low'],
            'CI_High': result['ci_high'],
            'Evaluation_Time_Seconds': result.get('evaluation_time', 0),
            'Load_Failed': result.get('model_load_failed', False),
            'Inference_Errors': result.get('inference_errors', 0),
            'Addition_Accuracy': result['by_operation']['+']['accuracy'] if '+' in result['by_operation'] else 0,
            'Subtraction_Accuracy': result['by_operation']['-']['accuracy'] if '-' in result['by_operation'] else 0,
            'Multiplication_Accuracy': result['by_operation']['*']['accuracy'] if '*' in result['by_operation'] else 0,
            'Division_Accuracy': result['by_operation']['/']['accuracy'] if '/' in result['by_operation'] else 0,
            'Exponentiation_Accuracy': result['by_operation']['**']['accuracy'] if '**' in result['by_operation'] else 0,
        })
    
    df = pd.DataFrame(summary_data)
    df.to_csv(csv_filename, index=False)
    logger.info(f"CSV summary saved to {csv_filename}")

def create_word_documentation(results: Dict[str, Dict[str, Any]]) -> None:
    """Create comprehensive Word documentation."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    doc_filename = os.path.join(args.output_dir, f"math_eval_report_{timestamp}.docx")
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('LLM Mathematical Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    # Calculate summary statistics
    successful_models = [name for name, result in results.items() if not result.get('model_load_failed', False)]
    failed_models = [name for name, result in results.items() if result.get('model_load_failed', False)]
    
    if successful_models:
        best_model = max(successful_models, key=lambda x: results[x]['accuracy'])
        worst_model = min(successful_models, key=lambda x: results[x]['accuracy'])
        avg_accuracy = np.mean([results[m]['accuracy'] for m in successful_models])
    else:
        best_model = worst_model = None
        avg_accuracy = 0
    
    summary_text = f"""
This report presents the evaluation results of {len(args.models)} language models on mathematical arithmetic tasks.

Key Findings:
• Total Models Evaluated: {len(args.models)}
• Successfully Loaded Models: {len(successful_models)}
• Failed to Load Models: {len(failed_models)}
• Questions per Model: {args.num_questions}
• Difficulty Level: {args.difficulty}
• Average Accuracy: {avg_accuracy:.1%}
"""
    
    if best_model:
        summary_text += f"• Best Performing Model: {best_model.split('/')[-1]} ({results[best_model]['accuracy']:.1%} accuracy)\n"
        summary_text += f"• Worst Performing Model: {worst_model.split('/')[-1]} ({results[worst_model]['accuracy']:.1%} accuracy)\n"
    
    summary_text += f"• Total Evaluation Time: {EVALUATION_STATS.get('total_duration', 0):.1f} seconds"
    
    doc.add_paragraph(summary_text)
    
    # Configuration Details
    doc.add_heading('Configuration', level=1)
    config_table = doc.add_table(rows=1, cols=2)
    config_table.style = 'Table Grid'
    hdr_cells = config_table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'
    
    config_items = [
        ('Number of Questions', args.num_questions),
        ('Max Tokens', args.max_tokens),
        ('Batch Size', args.batch_size),
        ('Timeout', args.timeout),
        ('Difficulty', args.difficulty),
        ('Device', DEVICE),
        ('Python Version', sys.version.split()[0]),
        ('PyTorch Version', torch.__version__),
        ('CUDA Available', torch.cuda.is_available())
    ]
    
    for param, value in config_items:
        row_cells = config_table.add_row().cells
        row_cells[0].text = str(param)
        row_cells[1].text = str(value)
    
    # Results Summary
    doc.add_heading('Results Summary', level=1)
    
    # Create results table
    results_table = doc.add_table(rows=1, cols=7)
    results_table.style = 'Table Grid'
    hdr_cells = results_table.rows[0].cells
    headers = ['Model', 'Accuracy', 'Correct', 'Incorrect', 'CI Low', 'CI High', 'Load Status']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    # Sort models by accuracy
    sorted_models = sorted(results.items(), key=lambda x: x[1]['accuracy'], reverse=True)
    
    for model_name, result in sorted_models:
        row_cells = results_table.add_row().cells
        row_cells[0].text = model_name.split('/')[-1]
        row_cells[1].text = f"{result['accuracy']:.1%}"
        row_cells[2].text = str(result['correct'])
        row_cells[3].text = str(result['incorrect'])
        row_cells[4].text = f"{result['ci_low']:.3f}"
        row_cells[5].text = f"{result['ci_high']:.3f}"
        row_cells[6].text = "Failed" if result.get('model_load_failed', False) else "Success"
    
    # Operation-specific Analysis
    doc.add_heading('Operation-Specific Performance', level=1)
    
    operations = ['+', '-', '*', '/']
    if args.difficulty in ['hard', 'mixed']:
        operations.append('**')
    
    for op in operations:
        op_name = {'+': 'Addition', '-': 'Subtraction', '*': 'Multiplication', 
                  '/': 'Division', '**': 'Exponentiation'}[op]
        
        doc.add_heading(f'{op_name} Performance', level=2)
        
        op_results = []
        for model_name, result in results.items():
            if not result.get('model_load_failed', False) and op in result['by_operation']:
                op_data = result['by_operation'][op]
                if op_data['total'] > 0:
                    op_results.append((model_name, op_data['accuracy'], op_data['correct'], op_data['total']))
        
        if op_results:
            op_results.sort(key=lambda x: x[1], reverse=True)
            
            op_table = doc.add_table(rows=1, cols=4)
            op_table.style = 'Table Grid'
            op_hdr = op_table.rows[0].cells
            op_hdr[0].text = 'Model'
            op_hdr[1].text = 'Accuracy'
            op_hdr[2].text = 'Correct'
            op_hdr[3].text = 'Total'
            
            for model_name, accuracy, correct, total in op_results:
                row_cells = op_table.add_row().cells
                row_cells[0].text = model_name.split('/')[-1]
                row_cells[1].text = f"{accuracy:.1%}"
                row_cells[2].text = str(correct)
                row_cells[3].text = str(total)
    
    # Error Analysis
    if MODEL_ERRORS:
        doc.add_heading('Error Analysis', level=1)
        
        for model_name, errors in MODEL_ERRORS.items():
            if errors:
                doc.add_heading(f'{model_name.split("/")[-1]} Errors', level=2)
                
                error_counts = {}
                for error in errors:
                    error_type = error['error_type']
                    error_counts[error_type] = error_counts.get(error_type, 0) + 1
                
                error_text = f"Total errors: {len(errors)}\n"
                for error_type, count in error_counts.items():
                    error_text += f"• {error_type}: {count}\n"
                
                doc.add_paragraph(error_text)
                
                # Show first few error details
                if len(errors) > 0:
                    doc.add_paragraph("Sample Errors:")
                    for i, error in enumerate(errors[:3]):  # Show first 3 errors
                        error_detail = f"{i+1}. {error['error_type']}: {error['error_message']}"
                        doc.add_paragraph(error_detail, style='List Bullet')
    
    # System Information
    doc.add_heading('System Information', level=1)
    sys_info = f"""
Evaluation Start Time: {EVALUATION_STATS.get('start_time', 'N/A')}
Evaluation End Time: {EVALUATION_STATS.get('end_time', 'N/A')}
Total Duration: {EVALUATION_STATS.get('total_duration', 0):.1f} seconds
Device Used: {DEVICE}
CUDA Available: {torch.cuda.is_available()}
Python Version: {sys.version}
PyTorch Version: {torch.__version__}
Total Questions Attempted: {EVALUATION_STATS.get('total_questions_attempted', 0)}
Total Questions Successful: {EVALUATION_STATS.get('total_questions_successful', 0)}
"""
    doc.add_paragraph(sys_info)
    
    # Save document
    doc.save(doc_filename)
    logger.info(f"Word documentation saved to {doc_filename}")

def plot_results(results: Dict[str, Dict[str, Any]]) -> None:
    """Create comprehensive visualizations of the results."""
    # Set the style
    sns.set(style="whitegrid")
    plt.rcParams.update({'font.size': 12})

    # 1. Overall accuracy bar chart with confidence intervals
    plt.figure(figsize=(16, 10))

    model_names = [name for name, result in results.items() if not result.get('model_load_failed', False)]
    if not model_names:
        logger.warning("No successful models to plot")
        return
        
    accuracies = [results[m]["accuracy"] * 100 for m in model_names]
    ci_lows = [results[m]["ci_low"] * 100 for m in model_names]
    ci_highs = [results[m]["ci_high"] * 100 for m in model_names]

    # Sort by accuracy
    sorted_indices = np.argsort(accuracies)[::-1]
    model_names = [model_names[i] for i in sorted_indices]
    accuracies = [accuracies[i] for i in sorted_indices]
    ci_lows = [ci_lows[i] for i in sorted_indices]
    ci_highs = [ci_highs[i] for i in sorted_indices]

    # Calculate error bar values
    yerr_low = [acc - low for acc, low in zip(accuracies, ci_lows)]
    yerr_high = [high - acc for acc, high in zip(accuracies, ci_highs)]

    # Clean up model names for display
    display_names = [name.split('/')[-1] for name in model_names]

    plt.bar(display_names, accuracies, color='skyblue', edgecolor='black')
    plt.errorbar(display_names, accuracies, yerr=[yerr_low, yerr_high], fmt='o', color='black', capsize=5)

    plt.xlabel('Model')
    plt.ylabel('Accuracy (%)')
    plt.title(f'Model Accuracy on Math Problems ({args.num_questions} questions each, {args.difficulty} difficulty)')
    plt.xticks(rotation=45, ha='right')
    plt.ylim(0, 100)

    for i, v in enumerate(accuracies):
        plt.text(i, v + 2, f"{v:.1f}%", ha='center')

    plt.tight_layout()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    plt.savefig(os.path.join(args.output_dir, f'overall_accuracy_{timestamp}.png'), dpi=300, bbox_inches='tight')

    # 2. Stacked bar chart showing correct vs incorrect
    plt.figure(figsize=(16, 10))

    model_names = [name for name, result in results.items() if not result.get('model_load_failed', False)]
    # Sort by accuracy
    sorted_indices = np.argsort([results[m]["accuracy"] for m in model_names])[::-1]
    model_names = [model_names[i] for i in sorted_indices]
    display_names = [name.split('/')[-1] for name in model_names]

    correct_vals = [results[m]["correct"] for m in model_names]
    incorrect_vals = [results[m]["incorrect"] for m in model_names]

    plt.bar(display_names, correct_vals, color='green', alpha=0.7, label='Correct')
    plt.bar(display_names, incorrect_vals, bottom=correct_vals, color='red', alpha=0.7, label='Incorrect')

    plt.xlabel('Model')
    plt.ylabel('Number of Questions')
    plt.title(f'Model Performance on Math Problems ({args.num_questions} questions each)')
    plt.xticks(rotation=45, ha='right')
    plt.legend()

    # Add accuracy percentages on top of bars
    total_height = [c + i for c, i in zip(correct_vals, incorrect_vals)]
    for i, (c, t) in enumerate(zip(correct_vals, total_height)):
        plt.text(i, t/2, f"{c/t*100:.1f}%", ha='center', va='center', color='white', fontweight='bold')

    plt.tight_layout()
    plt.savefig(os.path.join(args.output_dir, f'correct_vs_incorrect_{timestamp}.png'), dpi=300, bbox_inches='tight')

    # 3. Heatmap of operation-specific accuracy
    plt.figure(figsize=(16, 12))

    # Prepare data for heatmap
    operations = ['+', '-', '*', '/']
    if args.difficulty in ['hard', 'mixed']:
        operations.append('**')

    heatmap_data = []
    successful_models = [name for name, result in results.items() if not result.get('model_load_failed', False)]
    
    for model in successful_models:
        model_data = []
        for op in operations:
            if op in results[model]["by_operation"]:
                acc = results[model]["by_operation"][op]["accuracy"] * 100
                model_data.append(acc)
            else:
                model_data.append(0)
        heatmap_data.append(model_data)

    if heatmap_data:
        # Create DataFrame for heatmap
        display_names = [name.split('/')[-1] for name in successful_models]
        df = pd.DataFrame(heatmap_data, columns=operations, index=display_names)

        # Plot heatmap
        sns.heatmap(df, annot=True, cmap="YlGnBu", fmt=".1f", cbar_kws={'label': 'Accuracy (%)'})
        plt.title('Accuracy by Operation Type (%)')
        plt.tight_layout()
        plt.savefig(os.path.join(args.output_dir, f'operation_heatmap_{timestamp}.png'), dpi=300, bbox_inches='tight')

    # 4. Error distribution chart
    if MODEL_ERRORS:
        plt.figure(figsize=(14, 8))
        
        error_counts = {}
        for model_name, errors in MODEL_ERRORS.items():
            for error in errors:
                error_type = error['error_type']
                if error_type not in error_counts:
                    error_counts[error_type] = {}
                if model_name not in error_counts[error_type]:
                    error_counts[error_type][model_name] = 0
                error_counts[error_type][model_name] += 1
        
        if error_counts:
            # Create stacked bar chart for errors
            error_types = list(error_counts.keys())
            models_with_errors = list(set().union(*[list(error_counts[et].keys()) for et in error_types]))
            
            bottom = np.zeros(len(models_with_errors))
            colors = plt.cm.Set3(np.linspace(0, 1, len(error_types)))
            
            for i, error_type in enumerate(error_types):
                heights = [error_counts[error_type].get(model, 0) for model in models_with_errors]
                plt.bar(models_with_errors, heights, bottom=bottom, label=error_type, color=colors[i])
                bottom += heights
            
            plt.xlabel('Model')
            plt.ylabel('Number of Errors')
            plt.title('Error Distribution by Model and Type')
            plt.xticks(rotation=45, ha='right')
            plt.legend()
            plt.tight_layout()
            plt.savefig(os.path.join(args.output_dir, f'error_distribution_{timestamp}.png'), dpi=300, bbox_inches='tight')

    # 5. Evaluation time comparison
    plt.figure(figsize=(14, 8))
    
    models_with_time = [(name, result.get('evaluation_time', 0)) 
                       for name, result in results.items() 
                       if not result.get('model_load_failed', False) and result.get('evaluation_time', 0) > 0]
    
    if models_with_time:
        models_with_time.sort(key=lambda x: x[1])
        model_names_time = [name.split('/')[-1] for name, _ in models_with_time]
        eval_times = [time for _, time in models_with_time]
        
        plt.barh(model_names_time, eval_times, color='lightcoral')
        plt.xlabel('Evaluation Time (seconds)')
        plt.ylabel('Model')
        plt.title('Model Evaluation Time Comparison')
        
        for i, v in enumerate(eval_times):
            plt.text(v + 0.1, i, f"{v:.1f}s", va='center')
        
        plt.tight_layout()
        plt.savefig(os.path.join(args.output_dir, f'evaluation_times_{timestamp}.png'), dpi=300, bbox_inches='tight')

    # Show all plots
    plt.show()

def main() -> None:
    """Main evaluation function with comprehensive logging and documentation."""
    logger.info(f"Starting evaluation with {len(args.models)} models and {args.num_questions} questions each")
    logger.info(f"Difficulty level: {args.difficulty}")
    logger.info(f"Using device: {DEVICE}")
    
    # Log system information
    logger.info(f"System Info - Python: {sys.version.split()[0]}, PyTorch: {torch.__version__}, CUDA: {torch.cuda.is_available()}")

    start_time = time.time()
    
    try:
        results = evaluate_models()
        total_time = time.time() - start_time

        logger.info(f"Evaluation completed in {total_time:.2f} seconds")

        # Save comprehensive results
        save_results(results)

        # Create Word documentation
        try:
            create_word_documentation(results)
        except Exception as e:
            logger.error(f"Failed to create Word documentation: {str(e)}")
            logger.debug(traceback.format_exc())

        # Plot results
        try:
            plot_results(results)
        except Exception as e:
            logger.error(f"Failed to create plots: {str(e)}")
            logger.debug(traceback.format_exc())

        # Print summary to console
        print("\n" + "="*80)
        print("EVALUATION SUMMARY")
        print("="*80)
        
        successful_models = [name for name, result in results.items() if not result.get('model_load_failed', False)]
        failed_models = [name for name, result in results.items() if result.get('model_load_failed', False)]
        
        print(f"Total Models: {len(args.models)}")
        print(f"Successful Models: {len(successful_models)}")
        print(f"Failed Models: {len(failed_models)}")
        
        if successful_models:
            best_model = max(successful_models, key=lambda x: results[x]['accuracy'])
            worst_model = min(successful_models, key=lambda x: results[x]['accuracy'])
            avg_accuracy = np.mean([results[m]['accuracy'] for m in successful_models])
            
            print(f"Average Accuracy: {avg_accuracy:.1%}")
            print(f"Best Model: {best_model.split('/')[-1]} ({results[best_model]['accuracy']:.1%})")
            print(f"Worst Model: {worst_model.split('/')[-1]} ({results[worst_model]['accuracy']:.1%})")
        
        if failed_models:
            print(f"\nFailed Models:")
            for model in failed_models:
                print(f"  - {model.split('/')[-1]}")
        
        if MODEL_ERRORS:
            total_errors = sum(len(errors) for errors in MODEL_ERRORS.values())
            print(f"\nTotal Errors Encountered: {total_errors}")
        
        print(f"Total Evaluation Time: {total_time:.1f} seconds")
        print("="*80)

        return results
        
    except Exception as e:
        logger.error(f"Fatal error during evaluation: {str(e)}")
        logger.debug(traceback.format_exc())
        raise

# For notebook usage, provide a function to run with custom parameters
def run_evaluation(**kwargs):
    """Run evaluation with custom parameters (for use in notebooks)"""
    # Update configuration if parameters provided
    if kwargs:
        update_config(**kwargs)

    # Run evaluation
    return main()

if __name__ == "__main__":
    main()
